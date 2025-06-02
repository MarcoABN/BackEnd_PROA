import subprocess
from docx import Document
from tempfile import NamedTemporaryFile
from datetime import datetime
import os
import platform

def safe_str(value):
    return "" if value in [None, 0, "0", "0.0"] else str(value)

def safe_get(d, key):
    return d.get(key) if d.get(key) not in [None, 0, "0", "0.0"] else ""

def gerar_pdf_procuracao(cliente, embarcacao):
    doc = Document("docs/procuracao01OR.docx")

    print ('Entrou no utils', cliente)

    placeholders = {
        "${nomecliente}": safe_get(cliente, "nome"),
        "${nacionalidade}": safe_get(cliente, "nacionalidade"),
        "${enderecocompleto}": f"{safe_get(cliente, 'logradouro')}, {safe_get(cliente, 'numero')}, {safe_get(cliente, 'bairro')} - {safe_get(cliente, 'cidade')} - {safe_get(cliente, 'uf')}",
        "${identidade}": safe_get(cliente, "rg"),
        "${orgaoexpedidor}": safe_get(cliente, "orgEmissor"),
        "${cpfcliente}": safe_get(cliente, "cpfcnpj"),

        "${nomeemb}": safe_get(embarcacao, "nomeEmbarcacao"),
        "${numinscricao}": safe_get(embarcacao, "numInscricao"),
        "${tipo}": safe_get(embarcacao, "tipoEmbarcacao"),
        "${atividade}": safe_get(embarcacao, "tipoAtividade"),
        "${comprimento}": safe_str(embarcacao.get("compTotal")),
        "${tripulantes}": safe_str(embarcacao.get("qtdTripulantes")),
        "${boca}": safe_str(embarcacao.get("bocaMoldada")),
        "${passageiros}": safe_str(embarcacao.get("lotacao")),
        "${numcasco}": safe_get(embarcacao, "numCasco"),
        "${materialcasco}": safe_get(embarcacao, "matCasco"),
        "${potencia}": safe_str(embarcacao.get("potenciaMotor")),
        "${numserie}": safe_get(embarcacao, "numInscricao"),

        "${local}": safe_get(embarcacao, "cidade"),
        "${data}": datetime.now().strftime("%d/%m/%Y"),
    }

    # Substituir nos parágrafos
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Substituir nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Salvar documento temporário
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        tmp_docx_path = tmp_docx.name

    output_dir = os.path.dirname(tmp_docx_path)

    # Verifica o sistema operacional
    libreoffice_cmd = "soffice" if platform.system() == "Windows" else "libreoffice"

    # Executar conversão
    try:
        subprocess.run([
            libreoffice_cmd,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            tmp_docx_path
        ], check=True)
    except FileNotFoundError as e:
        raise RuntimeError(f"Erro: não foi possível encontrar '{libreoffice_cmd}'. Verifique se o LibreOffice está instalado e no PATH.") from e

    tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")

    # Ler PDF gerado
    with open(tmp_pdf_path, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()

    # Limpar arquivos temporários
    os.unlink(tmp_docx_path)
    os.unlink(tmp_pdf_path)
    print ('vai retornar o pdf')
    return pdf_bytes
